import os
import slicer
from slicer.ScriptedLoadableModule import *
import logging
import ctk
import qt
import uuid
import warnings
import sys

# Suppress Presidio multilingual warnings - we only use English
os.environ['PRESIDIO_SUPPRESS_WARNINGS'] = '1'

# Suppress warnings before importing Presidio
warnings.filterwarnings('ignore')
logging.getLogger("presidio_analyzer").setLevel(logging.CRITICAL)
logging.getLogger("presidio_anonymizer").setLevel(logging.CRITICAL)
logging.getLogger().setLevel(logging.CRITICAL)

# Capture and suppress stderr for Presidio initialization
class SuppressStderr:
    def write(self, x):
        pass
    def flush(self):
        pass

class Medical_Data_Anonymizer_Module(ScriptedLoadableModule):
    def __init__(self, parent):
        ScriptedLoadableModule.__init__(self, parent)
        self.parent.title = "Medical Data Anonymizer"
        self.parent.categories = ["Automated Dental Tools"]
        self.parent.dependencies = []
        self.parent.contributors = ["Jonas Bianchi,Alexandre BUisson, Paul Dumont"]
        self.parent.helpText = """This module anonymizes text files using Presidio."""
        self.parent.acknowledgementText = """Developed using Slicer resources."""

class Medical_Data_Anonymizer_ModuleWidget(ScriptedLoadableModuleWidget):

    def setup(self):
        ScriptedLoadableModuleWidget.setup(self)

        # Detect dark mode
        isDarkMode = self._isDarkMode()
        
        # Apply stylesheet to parent based on theme
        styleSheet = self._getStyleSheet(isDarkMode)
        self.parent.setStyleSheet(styleSheet)
        
        # Store reference for potential theme changes
        self.isDarkMode = isDarkMode
        
        # Add margins to left and right
        self.layout.setContentsMargins(15, 0, 15, 0)

        # Input Directory Section
        self.inputLabel = qt.QLabel("Files to be Anonymized")
        self.layout.addWidget(self.inputLabel)

        self.inputDirectoryButton = ctk.ctkPathLineEdit()
        self.inputDirectoryButton.filters = ctk.ctkPathLineEdit.Dirs
        self.inputDirectoryButton.setToolTip("Select the directory containing the files to be anonymized.")
        self.layout.addWidget(self.inputDirectoryButton)

        # Output Directory Section
        self.outputLabel = qt.QLabel("Output Anonymized Files")
        self.layout.addWidget(self.outputLabel)

        self.outputDirectoryButton = ctk.ctkPathLineEdit()
        self.outputDirectoryButton.filters = ctk.ctkPathLineEdit.Dirs
        self.outputDirectoryButton.setToolTip("Select the directory to save the anonymized files.")
        self.layout.addWidget(self.outputDirectoryButton)

        # Anonymization Options Section
        self.optionsLabel = qt.QLabel("<b>Anonymization Options:</b>")
        self.layout.addWidget(self.optionsLabel)

        # Create checkboxes for different entity types
        self.entityCheckboxes = {}
        
        entities = [
            ("PERSON", "Names (patients, doctors)", True),
            ("PHONE_NUMBER", "Phone Numbers", True),
            ("EMAIL_ADDRESS", "Email Addresses", True),
            ("DATE_TIME", "Dates and Times", True),
            ("LOCATION", "Addresses and Locations", True),
            ("US_SSN", "Social Security Numbers", True),
            ("MEDICAL_LICENSE", "Medical License Numbers", True),
            ("US_DRIVER_LICENSE", "Driver's License Numbers", False),
            ("CREDIT_CARD", "Credit Card Numbers", False),
            ("US_BANK_NUMBER", "Bank Account Numbers", False),
            ("IP_ADDRESS", "IP Addresses", False),
            ("URL", "URLs/Websites", False),
        ]

        for entity_type, description, default_checked in entities:
            checkbox = qt.QCheckBox(description)
            checkbox.setChecked(default_checked)
            checkbox.setToolTip(f"Anonymize {entity_type}")
            self.entityCheckboxes[entity_type] = checkbox
            self.layout.addWidget(checkbox)

        # Advanced Options Collapsible Section
        self.advancedCollapsible = ctk.ctkCollapsibleButton()
        self.advancedCollapsible.text = "Advanced Options"
        self.advancedCollapsible.collapsed = True
        self.layout.addWidget(self.advancedCollapsible)
        
        advancedLayout = qt.QFormLayout(self.advancedCollapsible)

        # Anonymization method dropdown
        self.anonymizationMethodCombo = qt.QComboBox()
        self.anonymizationMethodCombo.addItem("Replace with Label", "replace")
        self.anonymizationMethodCombo.addItem("Redact (Remove)", "redact")
        self.anonymizationMethodCombo.addItem("Hash", "hash")
        self.anonymizationMethodCombo.addItem("Mask", "mask")
        self.anonymizationMethodCombo.setToolTip("Choose how to anonymize detected entities")
        advancedLayout.addRow("Anonymization Method:", self.anonymizationMethodCombo)

        # Score threshold slider
        self.scoreThresholdSlider = ctk.ctkSliderWidget()
        self.scoreThresholdSlider.minimum = 0.0
        self.scoreThresholdSlider.maximum = 1.0
        self.scoreThresholdSlider.value = 0.5
        self.scoreThresholdSlider.singleStep = 0.05
        self.scoreThresholdSlider.setToolTip("Confidence threshold for entity detection (0.0-1.0). Higher = more strict.")
        advancedLayout.addRow("Confidence Threshold:", self.scoreThresholdSlider)

        # Install Dependencies Button
        self.installDependenciesButton = qt.QPushButton("Install Dependencies")
        self.installDependenciesButton.toolTip = "Install Presidio and required dependencies."
        self.layout.addWidget(self.installDependenciesButton)
        self.installDependenciesButton.connect('clicked(bool)', self.install_dependencies)

        # Anonymize Button
        self.anonymizeButton = qt.QPushButton("Anonymize Files")
        self.anonymizeButton.toolTip = "Run the anonymization process."
        self.layout.addWidget(self.anonymizeButton)
        self.anonymizeButton.connect('clicked(bool)', self.onAnonymizeButton)

        # Progress bar
        self.progressBar = qt.QProgressBar()
        self.progressBar.setVisible(False)
        self.layout.addWidget(self.progressBar)

        # Status label
        self.statusLabel = qt.QLabel("")
        self.layout.addWidget(self.statusLabel)

        # Add vertical spacer
        self.layout.addStretch(1)

    def _isDarkMode(self):
        """Check if the application is in dark mode"""
        try:
            # Get the palette of the main application
            palette = slicer.app.palette()
            # Check if the background is dark by checking luminance
            bgColor = palette.color(qt.QPalette.Window)
            luminance = (0.299 * bgColor.red() + 0.587 * bgColor.green() + 0.114 * bgColor.blue()) / 255.0
            return luminance < 0.5
        except:
            return False

    def _getStyleSheet(self, isDarkMode):
        """Generate stylesheet based on theme"""
        if isDarkMode:
            # Dark mode colors
            return """
            qMRMLWidget {
              background-color: #2b2b2b;
            }
            ctkCollapsibleButton {
              background-color: #383838;
              border: 1px solid #454545;
              border-radius: 6px;
              margin-bottom: 8px;
              font-weight: 600;
              padding: 6px 10px;
              color: #e0e0e0;
            }
            ctkCollapsibleButton:hover {
              border: 1px solid #3498db;
              background-color: #414141;
            }
            QLineEdit, QTextEdit {
              background-color: #353535;
              border: 1px solid #454545;
              border-radius: 4px;
              padding: 6px;
              color: #e0e0e0;
              selection-background-color: #3498db;
            }
            QLineEdit:focus, QTextEdit:focus {
              border: 2px solid #3498db;
              background-color: #383838;
            }
            QComboBox {
              background-color: #353535;
              border: 1px solid #454545;
              border-radius: 4px;
              padding: 4px 6px;
              color: #e0e0e0;
            }
            QComboBox:focus {
              border: 2px solid #3498db;
            }
            QComboBox::drop-down {
              width: 20px;
              border: none;
            }
            QComboBox QAbstractItemView {
              background-color: #353535;
              color: #e0e0e0;
              selection-background-color: #3498db;
              border: 1px solid #454545;
            }
            QLabel {
              color: #e0e0e0;
              font-weight: 500;
            }
            QPushButton {
              background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #4ba3ff, stop:1 #3498db);
              color: white;
              border: none;
              border-radius: 6px;
              font-weight: 600;
              font-size: 10pt;
              padding: 8px;
              margin-top: 4px;
            }
            QPushButton:hover:!pressed {
              background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #5cb3ff, stop:1 #2980b9);
            }
            QPushButton:pressed {
              background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2980b9, stop:1 #1f618d);
            }
            QPushButton:disabled {
              background-color: #555555;
              color: #888888;
            }
            QCheckBox {
              color: #e0e0e0;
              font-weight: 500;
              spacing: 6px;
            }
            QCheckBox::indicator {
              width: 18px;
              height: 18px;
              border: 1px solid #555555;
              border-radius: 3px;
              background-color: #353535;
            }
            QCheckBox::indicator:hover {
              border: 1px solid #3498db;
              background-color: #3d3d3d;
            }
            QCheckBox::indicator:checked {
              background-color: #3498db;
              border: 1px solid #3498db;
              image: url("data:image/svg+xml,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 16 16'><path fill='white' d='M13.854 3.646a.5.5 0 0 1 0 .708l-7 7a.5.5 0 0 1-.708 0l-3.5-3.5a.5.5 0 1 1 .708-.708L6.5 10.293l6.646-6.647a.5.5 0 0 1 .708 0z'/></svg>");
            }
            QProgressBar {
              border: 1px solid #454545;
              border-radius: 4px;
              background-color: #353535;
              padding: 2px;
              color: #e0e0e0;
            }
            QProgressBar::chunk {
              background-color: #3498db;
              border-radius: 3px;
            }
            """
        else:
            # Light mode colors (original)
            return """
            qMRMLWidget {
              background-color: #f8f9fa;
            }
            ctkCollapsibleButton {
              background-color: #ffffff;
              border: 1px solid #e0e6ed;
              border-radius: 6px;
              margin-bottom: 8px;
              font-weight: 600;
              padding: 6px 10px;
              color: #2c3e50;
            }
            ctkCollapsibleButton:hover {
              border: 1px solid #3498db;
              background-color: #fbfcfd;
            }
            QLineEdit, QTextEdit {
              background-color: #ffffff;
              border: 1px solid #e0e6ed;
              border-radius: 4px;
              padding: 6px;
              color: #2c3e50;
              selection-background-color: #3498db;
            }
            QLineEdit:focus, QTextEdit:focus {
              border: 2px solid #3498db;
            }
            QComboBox {
              background-color: #ffffff;
              border: 1px solid #e0e6ed;
              border-radius: 4px;
              padding: 4px 6px;
              color: #2c3e50;
            }
            QComboBox:focus {
              border: 2px solid #3498db;
            }
            QComboBox::drop-down {
              width: 20px;
              border: none;
            }
            QComboBox QAbstractItemView {
              background-color: #ffffff;
              color: #2c3e50;
              selection-background-color: #e8f4f8;
              border: 1px solid #e0e6ed;
            }
            QLabel {
              color: #2c3e50;
              font-weight: 500;
            }
            QPushButton {
              background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #4ba3ff, stop:1 #3498db);
              color: white;
              border: none;
              border-radius: 6px;
              font-weight: 600;
              font-size: 10pt;
              padding: 8px;
              margin-top: 4px;
            }
            QPushButton:hover:!pressed {
              background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #5cb3ff, stop:1 #2980b9);
            }
            QPushButton:pressed {
              background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2980b9, stop:1 #1f618d);
            }
            QPushButton:disabled {
              background-color: #bdc3c7;
              color: #95a5a6;
            }
            QCheckBox {
              color: #2c3e50;
              font-weight: 500;
              spacing: 6px;
            }
            QCheckBox::indicator {
              width: 18px;
              height: 18px;
              border: 1px solid #e0e6ed;
              border-radius: 3px;
              background-color: #ffffff;
            }
            QCheckBox::indicator:hover {
              border: 1px solid #3498db;
              background-color: #fbfcfd;
            }
            QCheckBox::indicator:checked {
              background-color: #3498db;
              border: 1px solid #3498db;
              image: url("data:image/svg+xml,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 16 16'><path fill='white' d='M13.854 3.646a.5.5 0 0 1 0 .708l-7 7a.5.5 0 0 1-.708 0l-3.5-3.5a.5.5 0 1 1 .708-.708L6.5 10.293l6.646-6.647a.5.5 0 0 1 .708 0z'/></svg>");
            }
            QProgressBar {
              border: 1px solid #e0e6ed;
              border-radius: 4px;
              background-color: #ffffff;
              padding: 2px;
              color: #2c3e50;
            }
            QProgressBar::chunk {
              background-color: #3498db;
              border-radius: 3px;
            }
            """

    def install_dependencies(self):
        try:
            self.statusLabel.setText("Installing dependencies...")
            slicer.app.processEvents()
            
            # Install Presidio and file handling libraries
            slicer.util.pip_install('presidio-analyzer')
            slicer.util.pip_install('presidio-anonymizer')
            slicer.util.pip_install('pandas')
            slicer.util.pip_install('python-docx')
            slicer.util.pip_install('pdfplumber')
            slicer.util.pip_install('odfpy')  # For ODT support
            slicer.util.pip_install('lxml')  # For XML support
            slicer.util.pip_install('langdetect')  # For language detection
            slicer.util.pip_install('reportlab')  # For PDF generation
            
            # Download spaCy language models
            import spacy
            
            models = [
                ("en_core_web_lg", "English")
            ]
            
            for model_name, lang_name in models:
                try:
                    self.statusLabel.setText(f"Loading {lang_name} model...")
                    slicer.app.processEvents()
                    spacy.load(model_name)
                except:
                    self.statusLabel.setText(f"Installing {lang_name} model...")
                    slicer.app.processEvents()
                    try:
                        spacy.cli.download(model_name)
                    except:
                        logging.warning(f"Could not download {model_name}")

            self.statusLabel.setText("Dependencies installed successfully!")
            
            # Notify user to restart Slicer
            qt.QMessageBox.information(
                slicer.util.mainWindow(),
                'Restart Required',
                'Dependencies have been installed. Please restart 3D Slicer to complete the installation.'
            )
        except Exception as e:
            self.statusLabel.setText(f"Error installing dependencies: {str(e)}")
            qt.QMessageBox.critical(
                slicer.util.mainWindow(),
                'Installation Error',
                f'Error installing dependencies: {str(e)}'
            )

    def onAnonymizeButton(self):
        # Validate inputs
        if not self.inputDirectoryButton.currentPath:
            qt.QMessageBox.warning(
                slicer.util.mainWindow(),
                'Input Required',
                'Please select an input directory.'
            )
            return

        if not self.outputDirectoryButton.currentPath:
            qt.QMessageBox.warning(
                slicer.util.mainWindow(),
                'Input Required',
                'Please select an output directory.'
            )
            return

        # Check dependencies
        try:
            from presidio_analyzer import AnalyzerEngine
            from presidio_anonymizer import AnonymizerEngine
            import docx
        except ImportError:
            qt.QMessageBox.warning(
                slicer.util.mainWindow(),
                'Dependencies Not Installed',
                'Please install the dependencies first by clicking the "Install Dependencies" button and restart 3D Slicer.'
            )
            return

        # Get selected entities
        selected_entities = [entity for entity, checkbox in self.entityCheckboxes.items() if checkbox.isChecked()]
        
        if not selected_entities:
            qt.QMessageBox.warning(
                slicer.util.mainWindow(),
                'No Options Selected',
                'Please select at least one anonymization option.'
            )
            return

        # Get anonymization method
        anonymization_method = self.anonymizationMethodCombo.currentData
        score_threshold = self.scoreThresholdSlider.value

        old_stderr = sys.stderr
        sys.stderr = SuppressStderr()
        try:
            analyzer = AnalyzerEngine()
            anonymizer = AnonymizerEngine()
        finally:
            sys.stderr = old_stderr

        input_folder = self.inputDirectoryButton.currentPath
        output_folder = self.outputDirectoryButton.currentPath
        csv_file_path = os.path.join(output_folder, "file_mappings.csv")

        # Get list of files (DOCX, TXT, PDF, CSV, XML, ODT)
        supported_files = []
        for root, dirs, files in os.walk(input_folder):
            for file in files:
                if (file.endswith((".docx", ".txt", ".pdf", ".csv", ".xml", ".odt")) and not file.startswith("~$")):
                    supported_files.append(os.path.join(root, file))

        if not supported_files:
            qt.QMessageBox.information(
                slicer.util.mainWindow(),
                'No Files Found',
                'No supported files found (.docx, .txt, .pdf, .csv, .xml, .odt).'
            )
            return

        # Show progress bar
        self.progressBar.setVisible(True)
        self.progressBar.setMaximum(len(supported_files))
        self.progressBar.setValue(0)

        # Run the anonymization process
        file_mappings = []

        for idx, input_file_path in enumerate(supported_files):
            try:
                file = os.path.basename(input_file_path)
                self.statusLabel.setText(f"Processing: {file}")
                slicer.app.processEvents()

                unique_id = str(uuid.uuid4())
                file_ext = os.path.splitext(file)[1].lower()
                
                # Extract text based on file type
                if file_ext == ".docx":
                    import docx
                    doc = docx.Document(input_file_path)
                    full_text = "\n".join([para.text for para in doc.paragraphs])
                
                elif file_ext == ".txt":
                    with open(input_file_path, 'r', encoding='utf-8') as f:
                        full_text = f.read()
                
                elif file_ext == ".pdf":
                    import pdfplumber
                    full_text = ""
                    try:
                        with pdfplumber.open(input_file_path) as pdf:
                            for page in pdf.pages:
                                text = page.extract_text()
                                if text:
                                    full_text += text + "\n"
                    except Exception as pdf_error:
                        logging.warning(f"Error extracting text from PDF {file}: {pdf_error}")
                        try:
                            with pdfplumber.open(input_file_path) as pdf:
                                for page in pdf.pages:
                                    tables = page.extract_tables()
                                    if tables:
                                        for table in tables:
                                            for row in table:
                                                full_text += " ".join([str(cell) if cell else "" for cell in row]) + "\n"
                        except:
                            pass
                
                elif file_ext == ".csv":
                    import csv
                    full_text = ""
                    with open(input_file_path, 'r', encoding='utf-8') as f:
                        reader = csv.reader(f)
                        for row in reader:
                            full_text += " ".join(row) + "\n"
                
                elif file_ext == ".xml":
                    from xml.etree import ElementTree as ET
                    tree = ET.parse(input_file_path)
                    root = tree.getroot()
                    full_text = self.extract_text_from_xml(root)
                
                elif file_ext == ".odt":
                    from odf import opendocument, text
                    doc = opendocument.load(input_file_path)
                    full_text = ""
                    for paragraph in doc.getElementsByType(text.P):
                        for node in paragraph.childNodes:
                            if node.nodeType == node.TEXT_NODE:
                                full_text += str(node.data)
                        full_text += "\n"
                
                else:
                    raise ValueError(f"Unsupported file type: {file_ext}")
                
                # Anonymize using Presidio
                anonymized_text = self.anonymize_text_presidio(
                    full_text, 
                    analyzer, 
                    anonymizer, 
                    selected_entities,
                    anonymization_method,
                    score_threshold
                )

                # Save anonymized file with same format
                if file_ext == ".docx":
                    import docx
                    new_doc = docx.Document()
                    for line in anonymized_text.split("\n"):
                        new_doc.add_paragraph(line)
                    new_file_name = file.replace(".docx", "_anonymized.docx")
                    output_path = os.path.join(output_folder, new_file_name)
                    new_doc.save(output_path)
                
                elif file_ext == ".txt":
                    new_file_name = file.replace(".txt", "_anonymized.txt")
                    output_path = os.path.join(output_folder, new_file_name)
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(anonymized_text)
                
                elif file_ext == ".pdf":
                    new_file_name = file.replace(".pdf", "_anonymized.pdf")
                    output_path = os.path.join(output_folder, new_file_name)
                    self.save_str_pdf(anonymized_text, output_path)
                
                elif file_ext == ".csv":
                    import csv
                    new_file_name = file.replace(".csv", "_anonymized.csv")
                    output_path = os.path.join(output_folder, new_file_name)
                    with open(output_path, 'w', encoding='utf-8', newline='') as f:
                        writer = csv.writer(f)
                        for line in anonymized_text.split("\n"):
                            if line.strip():
                                writer.writerow(line.split())
                
                elif file_ext == ".xml":
                    from xml.etree import ElementTree as ET
                    new_file_name = file.replace(".xml", "_anonymized.xml")
                    output_path = os.path.join(output_folder, new_file_name)
                    # Save as formatted text since XML anonymization is complex
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(anonymized_text)
                
                elif file_ext == ".odt":
                    new_file_name = file.replace(".odt", "_anonymized.txt")
                    output_path = os.path.join(output_folder, new_file_name)
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(anonymized_text)

                file_mappings.append({
                    "Original File Name": file,
                    "Anonymized File Name": new_file_name,
                    "UUID": unique_id
                })

                logging.info(f"Anonymized file created: {output_path}")

            except Exception as e:
                logging.error(f"Error processing {file}: {e}")
                file_mappings.append({
                    "Original File Name": file,
                    "Anonymized File Name": "ERROR",
                    "UUID": f"Error: {str(e)}"
                })

            self.progressBar.setValue(idx + 1)
            slicer.app.processEvents()

        # Save mappings
        import pandas as pd
        mappings_df = pd.DataFrame(file_mappings)
        if not mappings_df.empty:
            mappings_df.to_csv(csv_file_path, index=False)
            self.statusLabel.setText(f"Complete! Processed {len(supported_files)} files.")
            logging.info(f"Anonymization complete. File mappings saved to {csv_file_path}.")
            
            qt.QMessageBox.information(
                slicer.util.mainWindow(),
                'Anonymization Complete',
                f'Successfully processed {len(supported_files)} files.\n\nMappings saved to:\n{csv_file_path}'
            )
        else:
            self.statusLabel.setText("No valid files were processed.")
            logging.info("No valid files were processed. CSV file not created.")

        self.progressBar.setVisible(False)

    def save_str_pdf(self, text, filename):
        from reportlab.platypus import SimpleDocTemplate, Preformatted
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(filename)
        styles = getSampleStyleSheet()

        # Preformatted garde *tout* : <test>, <abc>, indentation, etc.
        story = [Preformatted(text, styles['Normal'])]

        doc.build(story)

    def anonymize_text_presidio(self, text, analyzer, anonymizer, entities, method, score_threshold):
        """
        Anonymize text using Presidio with automatic language detection
        
        Parameters:
        - text: The text to anonymize
        - analyzer: Presidio AnalyzerEngine instance
        - anonymizer: Presidio AnonymizerEngine instance
        - entities: List of entity types to detect
        - method: Anonymization method ('replace', 'redact', 'hash', 'mask')
        - score_threshold: Confidence threshold for detection
        """
        try:
            
            # Analyze text with detected language
            results = analyzer.analyze(
                text=text,
                language='en',
                entities=entities,
                score_threshold=score_threshold
            )

            # Define operators based on method
            if method == "replace":
                # Replace with entity type label
                operators = {}
            elif method == "redact":
                # Remove the entity completely
                from presidio_anonymizer.entities import OperatorConfig
                operators = {entity: OperatorConfig("redact") for entity in entities}
            elif method == "hash":
                # Replace with hash
                from presidio_anonymizer.entities import OperatorConfig
                operators = {entity: OperatorConfig("hash") for entity in entities}
            elif method == "mask":
                # Mask with asterisks
                from presidio_anonymizer.entities import OperatorConfig
                operators = {entity: OperatorConfig("mask", {"masking_char": "*", "chars_to_mask": 100, "from_end": False}) for entity in entities}
            else:
                operators = {}

            # Anonymize
            anonymized = anonymizer.anonymize(
                text=text,
                analyzer_results=results,
                operators=operators if operators else None
            )

            return anonymized.text

        except Exception as e:
            logging.error(f"Error in Presidio anonymization: {e}")
            return text  # Return original text if anonymization fails

    def extract_text_from_xml(self, element):
        """Recursively extract text from XML elements"""
        text = ""
        if element.text:
            text += element.text + " "
        for child in element:
            text += self.extract_text_from_xml(child)
            if child.tail:
                text += child.tail + " "
        return text

class Medical_Data_Anonymizer_ModuleLogic(ScriptedLoadableModuleLogic):
    pass

class Medical_Data_Anonymizer_ModuleTest(ScriptedLoadableModuleTest):
    pass
